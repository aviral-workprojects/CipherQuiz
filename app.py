"""
DocuSplit Pro - Complete Application
====================================
MVP + All Iterations (v5.0 - Production Ready)

Features:
- Two Modes: General Document Splitter & Quiz Generator
- Upload Word documents
- Split by text pattern
- Generate individual PDFs
- Encryption options:
  * No encryption
  * Same password for all
  * Different password per section
  * Sequential password flow (perfect for quizzes!)
- Combinations:
  * Random shuffle
  * Grouped shuffle (Quiz Mode)
  * Fixed positions
  * Multiple sets generation (5-30)
- Quiz Mode Exclusive:
  * Grouped shuffling by difficulty (Easy/Medium/Hard)
  * Automatic answer key generation
  * Verification checklists
  * Quiz-optimized settings
- Production Features (v5.0):
  * Comprehensive error handling
  * Input validation at every step
  * Progress indicators
  * Help tooltips
  * Example templates
  * Export/Import configurations
  * Deployment ready

Version: 5.0.0 - Production Ready
"""

import streamlit as st
import os
import re
import tempfile
import shutil
import zipfile
import itertools
import random
import json
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import List, Tuple, Optional, Dict

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from reportlab.lib.units import inch
from PyPDF2 import PdfReader, PdfWriter


# =============================================================================
# PAGE CONFIGURATION
# =============================================================================

st.set_page_config(
    page_title="DocuSplit Pro",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)


# =============================================================================
# DATA MODELS
# =============================================================================

@dataclass
class Section:
    """Represents a document section."""
    number: int
    title: str
    content: str
    original_number: int


# =============================================================================
# DOCUMENT PARSER
# =============================================================================

class DocumentParser:
    """Parses Word documents and extracts sections."""
    
    def __init__(self):
        self.document = None
        self.full_text = ""
    
    def get_preview_text(self, docx_path: str, max_chars: int = 1000) -> str:
        """Get preview text from document."""
        try:
            doc = Document(docx_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            if len(text) > max_chars:
                return text[:max_chars] + "..."
            return text
        except Exception as e:
            raise Exception(f"Error reading document: {str(e)}")
    
    def parse_document(self, docx_path: str, pattern: str) -> List[Section]:
        """Parse document and extract sections based on pattern."""
        try:
            doc = Document(docx_path)
            self.document = doc
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            self.full_text = "\n".join(paragraphs)
            sections = self._extract_sections(pattern)
            
            return sections
        except Exception as e:
            raise Exception(f"Error parsing document: {str(e)}")
    
    def _extract_sections(self, pattern: str) -> List[Section]:
        """Extract sections from full text using regex pattern."""
        sections = []
        
        try:
            regex = re.compile(pattern)
        except re.error as e:
            raise Exception(f"Invalid regex pattern: {str(e)}")
        
        matches = list(regex.finditer(self.full_text))
        
        if not matches:
            return []
        
        for i, match in enumerate(matches):
            section_start = match.start()
            section_end = matches[i + 1].start() if i + 1 < len(matches) else len(self.full_text)
            
            section_text = self.full_text[section_start:section_end].strip()
            lines = section_text.split('\n', 1)
            title = lines[0].strip()
            content = lines[1].strip() if len(lines) > 1 else ""
            
            section = Section(
                number=i + 1,
                title=title,
                content=content,
                original_number=i + 1
            )
            
            sections.append(section)
        
        return sections


# =============================================================================
# PDF GENERATOR
# =============================================================================

class PDFGenerator:
    """Generates PDF files from document sections."""
    
    def __init__(self):
        self.page_width, self.page_height = letter
        self.margin = 1 * inch
        self.content_width = self.page_width - (2 * self.margin)
    
    def create_pdf(self, section: Section, output_path: str, 
                   font: str = "Helvetica", font_size: int = 12):
        """Create a PDF from a section."""
        try:
            c = canvas.Canvas(output_path, pagesize=letter)
            y = self.page_height - self.margin
            
            # Draw title
            y = self._draw_title(c, section.title, y)
            y -= 0.3 * inch
            
            # Draw content
            y = self._draw_content(c, section.content, y, font, font_size)
            
            c.save()
        except Exception as e:
            raise Exception(f"Error creating PDF: {str(e)}")
    
    def _draw_title(self, c: canvas.Canvas, title: str, y: float) -> float:
        """Draw section title on PDF."""
        c.setFont("Helvetica-Bold", 16)
        c.setFillColorRGB(0.2, 0.2, 0.2)
        
        title_lines = simpleSplit(title, "Helvetica-Bold", 16, self.content_width)
        
        for line in title_lines:
            if y < self.margin:
                c.showPage()
                y = self.page_height - self.margin
                c.setFont("Helvetica-Bold", 16)
            
            c.drawString(self.margin, y, line)
            y -= 20
        
        return y
    
    def _draw_content(self, c: canvas.Canvas, content: str, y: float,
                     font: str, font_size: int) -> float:
        """Draw section content on PDF."""
        c.setFont(font, font_size)
        c.setFillColorRGB(0, 0, 0)
        
        paragraphs = content.split('\n')
        
        for para in paragraphs:
            if not para.strip():
                y -= font_size
                continue
            
            lines = simpleSplit(para, font, font_size, self.content_width)
            
            for line in lines:
                if y < self.margin + font_size:
                    c.showPage()
                    y = self.page_height - self.margin
                    c.setFont(font, font_size)
                
                c.drawString(self.margin, y, line)
                y -= font_size + 2
            
            y -= font_size * 0.5
        
        return y


# =============================================================================
# PDF ENCRYPTOR
# =============================================================================

class PDFEncryptor:
    """Handles PDF encryption."""
    
    @staticmethod
    def encrypt_pdf(input_path: str, output_path: str, password: str):
        """Encrypt a PDF with a password."""
        try:
            reader = PdfReader(input_path)
            writer = PdfWriter()
            
            for page in reader.pages:
                writer.add_page(page)
            
            writer.encrypt(password)
            
            with open(output_path, "wb") as f:
                writer.write(f)
        except Exception as e:
            raise Exception(f"Error encrypting PDF: {str(e)}")


# =============================================================================
# COMBINATION GENERATOR
# =============================================================================

class CombinationGenerator:
    """Generates multiple shuffled combinations of sections."""
    
    @staticmethod
    def generate_combinations(sections: List[Section], 
                            num_combinations: int,
                            shuffle_mode: str,
                            keep_first_fixed: bool,
                            keep_last_fixed: bool) -> List[List[Section]]:
        """
        Generate multiple shuffled combinations.
        
        Args:
            sections: List of Section objects
            num_combinations: Number of combinations to generate
            shuffle_mode: 'random' or 'none'
            keep_first_fixed: Keep first section in position 1
            keep_last_fixed: Keep last section in last position
            
        Returns:
            List of section combinations (each is a list of reordered sections)
        """
        if shuffle_mode == 'none':
            # No shuffling - return original order multiple times
            return [sections[:] for _ in range(num_combinations)]
        
        # Random shuffling
        combinations = []
        section_count = len(sections)
        
        # Determine which sections can be shuffled
        if keep_first_fixed and keep_last_fixed:
            # Only middle sections shuffle
            fixed_first = sections[0]
            fixed_last = sections[-1]
            shuffleable = sections[1:-1]
        elif keep_first_fixed:
            # First fixed, rest shuffle
            fixed_first = sections[0]
            fixed_last = None
            shuffleable = sections[1:]
        elif keep_last_fixed:
            # Last fixed, rest shuffle
            fixed_first = None
            fixed_last = sections[-1]
            shuffleable = sections[:-1]
        else:
            # All shuffle
            fixed_first = None
            fixed_last = None
            shuffleable = sections[:]
        
        # Generate unique combinations
        attempts = 0
        max_attempts = num_combinations * 10
        
        while len(combinations) < num_combinations and attempts < max_attempts:
            attempts += 1
            
            # Shuffle the shuffleable sections
            shuffled = shuffleable[:]
            random.shuffle(shuffled)
            
            # Rebuild full combination
            if keep_first_fixed and keep_last_fixed:
                combo = [fixed_first] + shuffled + [fixed_last]
            elif keep_first_fixed:
                combo = [fixed_first] + shuffled
            elif keep_last_fixed:
                combo = shuffled + [fixed_last]
            else:
                combo = shuffled
            
            # Update section numbers for this combination
            renumbered_combo = []
            for idx, section in enumerate(combo, start=1):
                new_section = Section(
                    number=idx,
                    title=section.title,
                    content=section.content,
                    original_number=section.original_number
                )
                renumbered_combo.append(new_section)
            
            # Check if this combination is unique
            combo_signature = tuple(s.original_number for s in renumbered_combo)
            existing_signatures = [tuple(s.original_number for s in c) for c in combinations]
            
            if combo_signature not in existing_signatures:
                combinations.append(renumbered_combo)
        
        # If we couldn't generate enough unique combinations, inform user
        if len(combinations) < num_combinations:
            # Fill remaining with original order
            while len(combinations) < num_combinations:
                combinations.append(sections[:])
        
        return combinations


# =============================================================================
# QUIZ GENERATOR
# =============================================================================

class QuizGenerator:
    """Handles quiz-specific functionality including grouped shuffling and answer keys."""
    
    @staticmethod
    def generate_grouped_combinations(sections: List[Section],
                                     num_combinations: int,
                                     groups: Dict[str, List[int]],
                                     shuffle_settings: Dict[str, bool]) -> List[List[Section]]:
        """
        Generate combinations with grouped shuffling (Easy/Medium/Hard).
        
        Args:
            sections: List of Section objects
            num_combinations: Number of combinations to generate
            groups: Dictionary defining groups (e.g., {'easy': [1,2,3], 'medium': [4,5,6], 'hard': [7,8,9]})
            shuffle_settings: Which groups to shuffle (e.g., {'easy': True, 'medium': True, 'hard': False})
            
        Returns:
            List of section combinations with grouped shuffling
        """
        combinations = []
        
        # Separate sections into groups
        grouped_sections = {}
        for group_name, group_indices in groups.items():
            grouped_sections[group_name] = [s for s in sections if s.original_number in group_indices]
        
        # Generate combinations
        attempts = 0
        max_attempts = num_combinations * 20
        
        while len(combinations) < num_combinations and attempts < max_attempts:
            attempts += 1
            
            # Shuffle or keep original order for each group
            combo_parts = []
            
            for group_name in ['easy', 'medium', 'hard']:  # Maintain order
                if group_name in grouped_sections:
                    group_sections = grouped_sections[group_name][:]
                    
                    if shuffle_settings.get(group_name, False):
                        random.shuffle(group_sections)
                    
                    combo_parts.extend(group_sections)
            
            # Renumber sections for this combination
            renumbered_combo = []
            for idx, section in enumerate(combo_parts, start=1):
                new_section = Section(
                    number=idx,
                    title=section.title,
                    content=section.content,
                    original_number=section.original_number
                )
                renumbered_combo.append(new_section)
            
            # Check uniqueness
            combo_signature = tuple(s.original_number for s in renumbered_combo)
            existing_signatures = [tuple(s.original_number for s in c) for c in combinations]
            
            if combo_signature not in existing_signatures:
                combinations.append(renumbered_combo)
        
        # Fill remaining if needed
        while len(combinations) < num_combinations:
            combinations.append(sections[:])
        
        return combinations
    
    @staticmethod
    def generate_answer_keys(combinations: List[List[Section]],
                            answers: Dict[int, str],
                            combo_prefix: str = "Combo") -> Dict:
        """
        Generate answer keys for all combinations.
        
        Args:
            combinations: List of combinations (each is a list of sections)
            answers: Dictionary mapping original question numbers to answers
            combo_prefix: Prefix for combination names
            
        Returns:
            Dictionary with answer key data
        """
        answer_keys = {}
        
        for idx, combo in enumerate(combinations, start=1):
            combo_name = f"{combo_prefix}_{idx:02d}"
            
            # Build answer sequence for this combination
            combo_answers = []
            for position, section in enumerate(combo, start=1):
                combo_answers.append({
                    'position': position,
                    'original_question': section.original_number,
                    'question_title': section.title,
                    'answer': answers.get(section.original_number, "NOT SET")
                })
            
            answer_keys[combo_name] = combo_answers
        
        return answer_keys
    
    @staticmethod
    def generate_verification_checklist(combinations: List[List[Section]],
                                        answers: Dict[int, str],
                                        encryption_mode: str,
                                        combo_prefix: str = "Combo") -> str:
        """
        Generate a verification checklist for testing all combinations.
        
        Args:
            combinations: List of combinations
            answers: Dictionary of answers
            encryption_mode: Type of encryption used
            combo_prefix: Prefix for combination names
            
        Returns:
            Formatted checklist string
        """
        checklist = []
        checklist.append("=" * 70)
        checklist.append("QUIZ VERIFICATION CHECKLIST")
        checklist.append("=" * 70)
        checklist.append("")
        checklist.append(f"Total Combinations: {len(combinations)}")
        checklist.append(f"Encryption Mode: {encryption_mode}")
        checklist.append("")
        checklist.append("Test each combination by following these steps:")
        checklist.append("")
        
        for idx, combo in enumerate(combinations, start=1):
            combo_name = f"{combo_prefix}_{idx:02d}"
            checklist.append(f"[ ] {combo_name}")
            checklist.append(f"    Question Order: {' → '.join([f'Q{s.original_number}' for s in combo])}")
            checklist.append("")
            
            for position, section in enumerate(combo, start=1):
                pdf_name = f"Section_{position}.pdf"
                
                if encryption_mode == "sequential":
                    if position == 1:
                        checklist.append(f"    [ ] {pdf_name} opens without password")
                    else:
                        prev_section = combo[position - 2]
                        prev_answer = answers.get(prev_section.original_number, "NOT SET")
                        checklist.append(f"    [ ] {pdf_name} opens with password: {prev_answer}")
                elif encryption_mode == "different":
                    answer = answers.get(section.original_number, "NOT SET")
                    checklist.append(f"    [ ] {pdf_name} (Q{section.original_number}) password: {answer}")
                elif encryption_mode == "same":
                    checklist.append(f"    [ ] {pdf_name} uses common password")
                else:
                    checklist.append(f"    [ ] {pdf_name} no password")
            
            checklist.append("")
        
        checklist.append("=" * 70)
        checklist.append("VERIFICATION COMPLETE")
        checklist.append("=" * 70)
        
        return "\n".join(checklist)


# =============================================================================
# CONFIGURATION MANAGER
# =============================================================================

class ConfigurationManager:
    """Handles saving and loading application configurations."""
    
    @staticmethod
    def export_config(session_state) -> str:
        """
        Export current configuration to JSON.
        
        Args:
            session_state: Streamlit session state
            
        Returns:
            JSON string of configuration
        """
        config = {
            'version': '5.0.0',
            'export_date': datetime.now().isoformat(),
            'mode': session_state.get('mode'),
            'encryption_mode': session_state.get('encryption_mode', 'none'),
            'combinations_enabled': session_state.get('combinations_enabled', False),
            'num_combinations': session_state.get('num_combinations', 5),
            'shuffle_mode': session_state.get('shuffle_mode', 'random'),
            'keep_first_fixed': session_state.get('keep_first_fixed', False),
            'keep_last_fixed': session_state.get('keep_last_fixed', False),
        }
        
        # Quiz mode specific
        if session_state.get('mode') == 'quiz':
            config['quiz_shuffle_settings'] = session_state.get('quiz_shuffle_settings', {})
            config['quiz_groups'] = session_state.get('quiz_groups', {})
        
        return json.dumps(config, indent=2)
    
    @staticmethod
    def import_config(config_json: str) -> Dict:
        """
        Import configuration from JSON.
        
        Args:
            config_json: JSON string of configuration
            
        Returns:
            Configuration dictionary
        """
        try:
            config = json.loads(config_json)
            
            # Validate version compatibility
            if not config.get('version'):
                raise ValueError("Invalid configuration file - missing version")
            
            return config
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {str(e)}")
    
    @staticmethod
    def apply_config(session_state, config: Dict):
        """
        Apply imported configuration to session state.
        
        Args:
            session_state: Streamlit session state
            config: Configuration dictionary
        """
        # Apply general settings
        for key in ['mode', 'encryption_mode', 'combinations_enabled', 
                    'num_combinations', 'shuffle_mode', 'keep_first_fixed', 'keep_last_fixed']:
            if key in config:
                session_state[key] = config[key]
        
        # Apply quiz settings if present
        if config.get('mode') == 'quiz':
            if 'quiz_shuffle_settings' in config:
                session_state['quiz_shuffle_settings'] = config['quiz_shuffle_settings']
            if 'quiz_groups' in config:
                session_state['quiz_groups'] = config['quiz_groups']


# =============================================================================
# ERROR HANDLER
# =============================================================================

class ErrorHandler:
    """Centralized error handling with user-friendly messages."""
    
    @staticmethod
    def handle_file_error(error: Exception, context: str = "file operation") -> str:
        """
        Handle file-related errors.
        
        Args:
            error: The exception that occurred
            context: Context where error occurred
            
        Returns:
            User-friendly error message
        """
        error_messages = {
            'PermissionError': f"Permission denied during {context}. Check file permissions.",
            'FileNotFoundError': f"File not found during {context}. File may have been moved or deleted.",
            'IsADirectoryError': f"Expected a file but found a directory during {context}.",
            'OSError': f"System error during {context}. Check disk space and permissions.",
        }
        
        error_type = type(error).__name__
        return error_messages.get(error_type, f"Unexpected error during {context}: {str(error)}")
    
    @staticmethod
    def handle_pdf_error(error: Exception) -> str:
        """Handle PDF generation errors."""
        if "encryption" in str(error).lower():
            return "PDF encryption failed. Check that passwords are valid and try again."
        elif "permission" in str(error).lower():
            return "Cannot write PDF file. Close any open PDFs and try again."
        else:
            return f"PDF generation error: {str(error)}"
    
    @staticmethod
    def handle_document_error(error: Exception) -> str:
        """Handle document parsing errors."""
        if "corrupt" in str(error).lower():
            return "Document appears to be corrupted. Try re-saving it in Word and upload again."
        elif "xml" in str(error).lower():
            return "Document format error. Ensure file is saved as .docx (not .doc)."
        else:
            return f"Document parsing error: {str(error)}"
    
    @staticmethod
    def safe_operation(operation, error_context: str = "operation"):
        """
        Wrapper for safe operation execution with error handling.
        
        Args:
            operation: Function to execute
            error_context: Context description for error messages
            
        Returns:
            (success: bool, result or error_message)
        """
        try:
            result = operation()
            return True, result
        except Exception as e:
            error_msg = ErrorHandler.handle_file_error(e, error_context)
            return False, error_msg


# =============================================================================
# VALIDATORS
# =============================================================================

def validate_docx_file(file) -> Tuple[bool, str]:
    """Validate uploaded Word document."""
    if file is None:
        return False, "No file uploaded"
    
    if not file.name.lower().endswith('.docx'):
        return False, "File must be a .docx Word document"
    
    max_size = 50 * 1024 * 1024
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    
    if file_size > max_size:
        return False, f"File too large ({file_size / 1024 / 1024:.1f}MB). Maximum is 50MB"
    
    if file_size == 0:
        return False, "File is empty"
    
    return True, ""


def validate_pattern(pattern: str) -> Tuple[bool, str]:
    """Validate regex pattern."""
    if not pattern or not pattern.strip():
        return False, "Pattern cannot be empty"
    
    try:
        re.compile(pattern)
        return True, ""
    except re.error as e:
        return False, f"Invalid regex: {str(e)}"


def validate_passwords(passwords: Dict[int, str], mode: str) -> Tuple[bool, str]:
    """Validate password configuration."""
    if mode == "none":
        return True, ""
    
    empty_passwords = []
    for section_num, password in passwords.items():
        if not password or not password.strip():
            empty_passwords.append(section_num)
    
    if empty_passwords:
        return False, f"Missing passwords for sections: {', '.join(map(str, empty_passwords))}"
    
    return True, ""


# =============================================================================
# HELPERS
# =============================================================================

def cleanup_temp_files(temp_dir: str):
    """Clean up temporary files and directories."""
    try:
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
    except Exception as e:
        print(f"Warning: Could not cleanup temp files: {str(e)}")


def create_download_zip(file_paths: List[str], zip_name: str) -> str:
    """Create a ZIP file from list of file paths."""
    temp_dir = tempfile.gettempdir()
    zip_path = os.path.join(temp_dir, zip_name)
    
    try:
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for file_path in file_paths:
                if os.path.exists(file_path):
                    filename = os.path.basename(file_path)
                    zip_file.write(file_path, filename)
        
        return zip_path
    except Exception as e:
        raise Exception(f"Error creating ZIP file: {str(e)}")


# =============================================================================
# SESSION STATE
# =============================================================================

def init_session_state():
    """Initialize session state variables."""
    defaults = {
        'mode': None,  # 'general' or 'quiz'
        'document': None,
        'sections': None,
        'generated_pdfs': None,
        'encryption_mode': 'none',
        'passwords': {},
        'current_step': 1,
        'combinations_enabled': False,
        'num_combinations': 5,
        'shuffle_mode': 'random',
        'keep_first_fixed': False,
        'keep_last_fixed': False,
        'all_combinations': None,
        # Quiz mode specific
        'quiz_groups': {'easy': [1, 2, 3], 'medium': [4, 5, 6], 'hard': [7, 8, 9]},
        'quiz_shuffle_settings': {'easy': True, 'medium': True, 'hard': True},
        'answer_keys': None,
        'verification_checklist': None
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


# =============================================================================
# MAIN APPLICATION
# =============================================================================

def main():
    """Main application entry point."""
    
    init_session_state()
    
    # Header
    st.title("📄 DocuSplit Pro")
    st.markdown("**Split Word documents into individual PDFs with optional encryption**")
    st.markdown("---")
    
    # ==========================================================================
    # MODE SELECTION
    # ==========================================================================
    
    if st.session_state.mode is None:
        st.header("Welcome! Choose Your Mode")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📄 General Mode")
            st.markdown("""
            **For any document:**
            - Reports, chapters, sections
            - Training materials
            - Any document splitting needs
            
            **Features:**
            - Flexible patterns
            - Random shuffle
            - Fixed positions
            """)
            
            if st.button("📄 Use General Mode", use_container_width=True, type="primary"):
                st.session_state.mode = 'general'
                st.rerun()
        
        with col2:
            st.subheader("🎓 Quiz Mode")
            st.markdown("""
            **For quiz/competition:**
            - Exam questions
            - Treasure hunts
            - Sequential puzzles
            
            **Features:**
            - Grouped shuffling (Easy/Medium/Hard)
            - Answer key generation
            - Verification checklists
            """)
            
            if st.button("🎓 Use Quiz Mode", use_container_width=True, type="primary"):
                st.session_state.mode = 'quiz'
                st.rerun()
        
        st.markdown("---")
        st.info("💡 **Tip:** Quiz Mode is optimized for sequential puzzle competitions with answer-based passwords!")
        st.stop()
    
    # Show current mode with option to change
    mode_display = "🎓 Quiz Mode" if st.session_state.mode == 'quiz' else "📄 General Mode"
    
    col1, col2 = st.columns([4, 1])
    with col1:
        st.caption(f"**Current Mode:** {mode_display}")
    with col2:
        if st.button("🔄 Change Mode", type="secondary"):
            # Reset everything
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()
    
    st.markdown("---")
    
    # Sidebar
    with st.sidebar:
        st.header("ℹ️ About DocuSplit Pro")
        st.caption("Version 5.0.0 - Production Ready")
        
        st.markdown("""
        **Core Features:**
        - ✅ Upload Word documents (.docx)
        - ✅ Pattern-based splitting
        - ✅ Professional PDF generation
        - ✅ 4 encryption modes
        - ✅ Multiple combinations (5-30)
        - ✅ Batch downloads
        
        **Quiz Mode Extras:**
        - 🎓 Grouped shuffling
        - 📝 Answer key generation
        - ✅ Verification checklists
        """)
        
        st.markdown("---")
        
        st.header("📊 Current Session")
        
        # Show mode
        if st.session_state.mode:
            mode_icon = "🎓" if st.session_state.mode == 'quiz' else "📄"
            st.metric("Mode", f"{mode_icon} {st.session_state.mode.title()}")
        else:
            st.metric("Mode", "—")
        
        if st.session_state.sections:
            st.metric("Sections", len(st.session_state.sections))
            st.metric("Encryption", st.session_state.encryption_mode.title())
            if st.session_state.combinations_enabled:
                st.metric("Combinations", st.session_state.num_combinations)
            else:
                st.metric("Combinations", "Single")
        else:
            st.metric("Sections", "—")
            st.metric("Encryption", "—")
            st.metric("Combinations", "—")
        
        st.markdown("---")
        
        # Configuration Management
        st.header("⚙️ Configuration")
        
        with st.expander("💾 Save/Load Settings"):
            st.caption("Save your current settings for reuse")
            
            # Export configuration
            if st.button("📤 Export Config", use_container_width=True):
                config_json = ConfigurationManager.export_config(st.session_state)
                st.download_button(
                    label="⬇️ Download Config",
                    data=config_json,
                    file_name=f"docusplit_config_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    use_container_width=True
                )
            
            # Import configuration
            st.caption("Or load saved settings:")
            config_file = st.file_uploader(
                "Upload config file",
                type=['json'],
                key="config_upload",
                label_visibility="collapsed"
            )
            
            if config_file:
                try:
                    config_json = config_file.read().decode('utf-8')
                    config = ConfigurationManager.import_config(config_json)
                    
                    if st.button("📥 Apply Config", use_container_width=True):
                        ConfigurationManager.apply_config(st.session_state, config)
                        st.success("✅ Configuration applied!")
                        st.rerun()
                except Exception as e:
                    st.error(f"❌ Error loading config: {str(e)}")
        
        st.markdown("---")
        
        # Quick Tips
        with st.expander("💡 Quick Tips"):
            if st.session_state.mode == 'quiz':
                st.markdown("""
                **Quiz Mode Tips:**
                - Use Sequential encryption for answer-based passwords
                - Test one combination before distributing all
                - Keep answer keys secure
                - Use verification checklist systematically
                """)
            else:
                st.markdown("""
                **General Mode Tips:**
                - Use preset patterns when possible
                - Preview sections before generating
                - Enable combinations for variety
                - Download master ZIP for all files
                """)
        
        # Help & Support
        with st.expander("❓ Need Help?"):
            st.markdown("""
            **Common Issues:**
            
            **No sections found:**
            - Try different pattern
            - Check document formatting
            
            **PDFs won't open:**
            - Verify password is correct
            - Check case sensitivity
            
            **Generation failed:**
            - Close open PDF files
            - Check disk space
            - Try smaller batch first
            """)
        
        st.markdown("---")
        st.caption("© 2026 DocuSplit Pro")
        st.caption("All-in-One Document Splitter")
    
    # ==========================================================================
    # STEP 1: UPLOAD
    # ==========================================================================
    
    st.header("Step 1: Upload Document")
    
    uploaded_file = st.file_uploader(
        "Choose a Word document (.docx)",
        type=['docx'],
        help="Upload a .docx file to split"
    )
    
    if uploaded_file is not None:
        is_valid, error_msg = validate_docx_file(uploaded_file)
        
        if not is_valid:
            st.error(f"❌ {error_msg}")
            st.stop()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        st.session_state.document = tmp_path
        st.success(f"✅ Uploaded: {uploaded_file.name}")
        
        with st.expander("📄 Document Preview"):
            try:
                parser = DocumentParser()
                preview = parser.get_preview_text(tmp_path, 1000)
                st.text_area("First 1000 characters:", preview, height=200, disabled=True)
            except Exception as e:
                st.warning(f"Preview unavailable: {str(e)}")
    else:
        st.info("👆 Upload a Word document to begin")
        st.stop()
    
    st.markdown("---")
    
    # ==========================================================================
    # STEP 2: CONFIGURE SPLITTING
    # ==========================================================================
    
    st.header("Step 2: Configure Splitting")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("Splitting Pattern")
        
        pattern_type = st.radio(
            "How to identify sections:",
            ["Preset Pattern", "Custom Text", "Custom Regex"],
            help="Select section detection method"
        )
        
        if pattern_type == "Preset Pattern":
            preset = st.selectbox(
                "Select preset:",
                [
                    "Question 1, Question 2, ...",
                    "Q1, Q2, Q3, ...",
                    "Section 1, Section 2, ...",
                    "Chapter 1, Chapter 2, ...",
                    "Part 1, Part 2, ...",
                ]
            )
            
            preset_patterns = {
                "Question 1, Question 2, ...": r"(?i)Question\s+\d+",
                "Q1, Q2, Q3, ...": r"(?i)Q\d+",
                "Section 1, Section 2, ...": r"(?i)Section\s+\d+",
                "Chapter 1, Chapter 2, ...": r"(?i)Chapter\s+\d+",
                "Part 1, Part 2, ...": r"(?i)Part\s+\d+",
            }
            
            pattern = preset_patterns[preset]
            st.info(f"🔍 Pattern: `{pattern}`")
        
        elif pattern_type == "Custom Text":
            custom_text = st.text_input(
                "Enter text that marks sections:",
                placeholder="e.g., 'Question', 'Task'",
                help="Text followed by numbers"
            )
            
            if custom_text:
                pattern = rf"(?i){custom_text}\s+\d+"
                st.info(f"🔍 Pattern: `{pattern}`")
            else:
                pattern = None
                st.warning("⚠️ Enter a text pattern")
        
        else:
            pattern = st.text_input(
                "Enter regex pattern:",
                placeholder=r"(?i)Question\s+\d+",
                help="Advanced regex pattern"
            )
            
            if pattern:
                is_valid, error = validate_pattern(pattern)
                if not is_valid:
                    st.error(f"❌ {error}")
                    pattern = None
    
    with col2:
        st.subheader("Pattern Help")
        st.markdown("""
        **Examples:**
        
        `Question 1`
        `Question 2`
        → Preset
        
        `Task A`
        `Task B`
        → Custom: "Task"
        """)
    
    if pattern:
        try:
            parser = DocumentParser()
            sections = parser.parse_document(st.session_state.document, pattern)
            
            if sections:
                st.session_state.sections = sections
                st.success(f"✅ Found {len(sections)} section(s)")
                
                with st.expander("📋 Detected Sections"):
                    for i, section in enumerate(sections, 1):
                        st.markdown(f"**Section {i}:** {section.title}")
                        preview = section.content[:150] + "..." if len(section.content) > 150 else section.content
                        st.text(preview)
                        if i < len(sections):
                            st.markdown("---")
            else:
                st.warning("⚠️ No sections found. Try different pattern.")
                st.stop()
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            st.stop()
    else:
        st.warning("⚠️ Configure a pattern")
        st.stop()
    
    st.markdown("---")
    
    # ==========================================================================
    # STEP 2.5: QUIZ MODE SETTINGS (Quiz Mode Only)
    # ==========================================================================
    
    if st.session_state.mode == 'quiz':
        st.header("Step 2.5: Quiz Settings")
        
        # Validate section count for quiz mode
        if len(st.session_state.sections) != 9:
            st.warning(f"⚠️ Quiz Mode works best with exactly 9 questions. You have {len(st.session_state.sections)} sections.")
            st.info("💡 You can still continue, but grouped shuffling assumes 3 groups of 3 questions.")
        
        # Answer Entry
        st.subheader("📝 Enter Answers")
        st.markdown("**Enter the correct answer for each question (these can be used as passwords in Sequential mode)**")
        
        show_answers = st.checkbox("👁️ Show answers as you type", value=False)
        answer_type = "default" if show_answers else "password"
        
        col1, col2, col3 = st.columns(3)
        
        quiz_answers = {}
        
        for idx, section in enumerate(st.session_state.sections):
            col = col1 if idx % 3 == 0 else (col2 if idx % 3 == 1 else col3)
            
            with col:
                answer = st.text_input(
                    f"Q{section.original_number} Answer:",
                    type=answer_type,
                    key=f"quiz_ans_{section.original_number}",
                    help=f"Answer for: {section.title[:30]}..."
                )
                quiz_answers[section.original_number] = answer
        
        # Validate all filled
        if any(not ans for ans in quiz_answers.values()):
            st.warning("⚠️ Please fill all answer fields before proceeding")
            st.stop()
        
        st.session_state.passwords = quiz_answers  # Store for later use
        
        st.markdown("---")
        
        # Grouped Shuffle Settings
        st.subheader("🔀 Grouped Shuffle Settings")
        st.markdown("**Configure which question groups should be shuffled**")
        
        col1, col2, col3 = st.columns(3)
        
        # Auto-detect groups based on section count
        total_sections = len(st.session_state.sections)
        group_size = total_sections // 3
        
        easy_group = list(range(1, group_size + 1))
        medium_group = list(range(group_size + 1, 2 * group_size + 1))
        hard_group = list(range(2 * group_size + 1, total_sections + 1))
        
        st.session_state.quiz_groups = {
            'easy': easy_group,
            'medium': medium_group,
            'hard': hard_group
        }
        
        with col1:
            st.markdown("**🟢 Easy Questions**")
            st.info(f"Questions: {', '.join(map(str, easy_group))}")
            shuffle_easy = st.checkbox(
                "Shuffle Easy questions",
                value=True,
                help="Randomly reorder questions within Easy group"
            )
        
        with col2:
            st.markdown("**🟡 Medium Questions**")
            st.info(f"Questions: {', '.join(map(str, medium_group))}")
            shuffle_medium = st.checkbox(
                "Shuffle Medium questions",
                value=True,
                help="Randomly reorder questions within Medium group"
            )
        
        with col3:
            st.markdown("**🔴 Hard Questions**")
            st.info(f"Questions: {', '.join(map(str, hard_group))}")
            shuffle_hard = st.checkbox(
                "Shuffle Hard questions",
                value=True,
                help="Randomly reorder questions within Hard group"
            )
        
        st.session_state.quiz_shuffle_settings = {
            'easy': shuffle_easy,
            'medium': shuffle_medium,
            'hard': shuffle_hard
        }
        
        # Show example shuffle
        with st.expander("📋 Preview: How Shuffling Works"):
            st.markdown("**Example question orders with current settings:**")
            
            if shuffle_easy and shuffle_medium and shuffle_hard:
                st.success("✅ All groups shuffling - maximum variety!")
                st.markdown("Example: `Q3, Q1, Q2, Q6, Q4, Q5, Q9, Q7, Q8`")
            elif not (shuffle_easy or shuffle_medium or shuffle_hard):
                st.info("ℹ️ No shuffling - all combinations will have same order")
                st.markdown(f"All combos: `Q{', Q'.join(map(str, range(1, total_sections + 1)))}`")
            else:
                shuffling = []
                if shuffle_easy:
                    shuffling.append("Easy")
                if shuffle_medium:
                    shuffling.append("Medium")
                if shuffle_hard:
                    shuffling.append("Hard")
                st.info(f"ℹ️ Shuffling: {', '.join(shuffling)} groups only")
        
        st.markdown("---")
    
    # ==========================================================================
    # STEP 3: ENCRYPTION SETTINGS
    # ==========================================================================
    
    st.header("Step 3: Encryption Settings")
    
    encryption_mode = st.radio(
        "Choose encryption mode:",
        [
            "🔓 No Encryption",
            "🔐 Same Password for All",
            "🔑 Different Password per Section",
            "🔗 Sequential Password Flow"
        ],
        help="Select how to protect PDFs"
    )
    
    # Map display text to internal mode
    mode_map = {
        "🔓 No Encryption": "none",
        "🔐 Same Password for All": "same",
        "🔑 Different Password per Section": "different",
        "🔗 Sequential Password Flow": "sequential"
    }
    
    st.session_state.encryption_mode = mode_map[encryption_mode]
    
    passwords = {}
    
    if st.session_state.encryption_mode == "same":
        st.info("ℹ️ All PDFs will use the same password")
        
        single_password = st.text_input(
            "Enter password for all sections:",
            type="password",
            help="This password will protect all PDFs"
        )
        
        if single_password:
            for section in st.session_state.sections:
                passwords[section.number] = single_password
        else:
            st.warning("⚠️ Enter a password")
            st.stop()
    
    elif st.session_state.encryption_mode == "different":
        st.info("ℹ️ Each PDF will have its own password")
        
        show_passwords = st.checkbox("Show passwords", value=False)
        password_type = "default" if show_passwords else "password"
        
        col1, col2 = st.columns(2)
        
        for idx, section in enumerate(st.session_state.sections):
            with col1 if idx % 2 == 0 else col2:
                pwd = st.text_input(
                    f"Password for Section {section.number}:",
                    type=password_type,
                    key=f"pwd_{section.number}",
                    help=f"Password for {section.title}"
                )
                passwords[section.number] = pwd
        
        # Validate all filled
        if any(not pwd for pwd in passwords.values()):
            st.warning("⚠️ Fill all password fields")
            st.stop()
    
    elif st.session_state.encryption_mode == "sequential":
        st.info("ℹ️ Section 1 has no password. Each subsequent section requires a password.")
        st.markdown("""
        **How it works:**
        - Section 1: No password (accessible to everyone)
        - Section 2: Requires password you set
        - Section 3: Requires different password
        - ... and so on
        
        💡 Tip: In quiz mode, each answer becomes the next password!
        """)
        
        show_passwords = st.checkbox("Show passwords", value=False)
        password_type = "default" if show_passwords else "password"
        
        # First section - no password
        passwords[1] = None
        st.success("✅ Section 1: No password (first section always accessible)")
        
        col1, col2 = st.columns(2)
        
        for idx, section in enumerate(st.session_state.sections[1:], start=2):
            with col1 if idx % 2 == 0 else col2:
                pwd = st.text_input(
                    f"Password for Section {section.number}:",
                    type=password_type,
                    key=f"pwd_seq_{section.number}",
                    help=f"Required to open {section.title}"
                )
                passwords[section.number] = pwd
        
        # Validate all filled (except first)
        if any(not pwd for num, pwd in passwords.items() if num > 1):
            st.warning("⚠️ Fill all password fields (except Section 1)")
            st.stop()
    
    elif st.session_state.encryption_mode == "none":
        st.info("ℹ️ PDFs will not be password-protected")
        for section in st.session_state.sections:
            passwords[section.number] = None
    
    st.session_state.passwords = passwords
    
    st.markdown("---")
    
    # ==========================================================================
    # STEP 3.5: COMBINATIONS (OPTIONAL)
    # ==========================================================================
    
    st.header("Step 3.5: Combinations (Optional)")
    
    combinations_enabled = st.checkbox(
        "🔄 Generate Multiple Combinations",
        value=False,
        help="Create multiple shuffled versions of the document"
    )
    
    st.session_state.combinations_enabled = combinations_enabled
    
    if combinations_enabled:
        st.info("ℹ️ Generate multiple shuffled sets - perfect for distributing different versions!")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Shuffle Settings")
            
            if st.session_state.mode == 'quiz':
                st.info("🎓 **Quiz Mode:** Using grouped shuffle settings from Step 2.5")
                st.session_state.shuffle_mode = "grouped"  # Special mode for quiz
                
                # Show which groups are shuffling
                shuffling_groups = []
                for group_name, is_shuffling in st.session_state.quiz_shuffle_settings.items():
                    if is_shuffling:
                        shuffling_groups.append(group_name.title())
                
                if shuffling_groups:
                    st.success(f"✅ Shuffling: {', '.join(shuffling_groups)} groups")
                else:
                    st.warning("⚠️ No groups set to shuffle - all combos will be identical")
            else:
                # General mode - original shuffle options
                shuffle_mode = st.radio(
                    "Shuffle mode:",
                    [
                        "🔀 Random Shuffle",
                        "📋 No Shuffle (Same Order)"
                    ],
                    help="How to reorder sections"
                )
                
                st.session_state.shuffle_mode = "random" if "Random" in shuffle_mode else "none"
                
                if st.session_state.shuffle_mode == "random":
                    st.success("✅ Sections will be randomly shuffled")
                else:
                    st.info("ℹ️ All combinations will have same order")
        
        with col2:
            st.subheader("Fixed Positions")
            
            if st.session_state.mode == 'quiz':
                st.info("🎓 **Quiz Mode:** Groups maintain their position")
                st.markdown("""
                - Easy group: Always first
                - Medium group: Always middle
                - Hard group: Always last
                - Shuffling happens *within* each group
                """)
            else:
                # General mode - original fixed position options
                keep_first = st.checkbox(
                    "📌 Keep first section in position 1",
                    value=False,
                    help="First section always stays first"
                )
                
                keep_last = st.checkbox(
                    "📌 Keep last section in last position",
                    value=False,
                    help="Last section always stays last"
                )
                
                st.session_state.keep_first_fixed = keep_first
                st.session_state.keep_last_fixed = keep_last
                
                if keep_first and keep_last:
                    st.info(f"ℹ️ Only middle {len(st.session_state.sections) - 2} sections will shuffle")
                elif keep_first:
                    st.info(f"ℹ️ Section 1 fixed, {len(st.session_state.sections) - 1} sections shuffle")
                elif keep_last:
                    st.info(f"ℹ️ Last section fixed, {len(st.session_state.sections) - 1} sections shuffle")
        
        st.markdown("---")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Number of Combinations")
            
            num_combinations = st.selectbox(
                "How many combinations to generate:",
                [5, 10, 15, 20, 25, 30],
                index=0,
                help="Each combination will be a separate folder"
            )
            
            st.session_state.num_combinations = num_combinations
        
        with col2:
            st.subheader("Preview")
            st.markdown(f"**Will generate:**")
            st.metric("Total Folders", num_combinations)
            st.metric("PDFs per Folder", len(st.session_state.sections))
            st.metric("Total PDFs", num_combinations * len(st.session_state.sections))
        
        # Show example combinations
        with st.expander("📋 Example Combinations Preview"):
            st.markdown("**Sample of how sections might be ordered:**")
            
            # Generate sample preview (just 3 examples)
            preview_combos = CombinationGenerator.generate_combinations(
                st.session_state.sections,
                min(3, num_combinations),
                st.session_state.shuffle_mode,
                st.session_state.keep_first_fixed,
                st.session_state.keep_last_fixed
            )
            
            for idx, combo in enumerate(preview_combos, start=1):
                st.markdown(f"**Combo {idx}:** {', '.join([f'Q{s.original_number}' for s in combo])}")
            
            if num_combinations > 3:
                st.markdown(f"... and {num_combinations - 3} more combinations")
        
        # Encryption compatibility warning
        if st.session_state.encryption_mode == "sequential":
            st.warning("⚠️ **Sequential encryption with combinations:** Each combination will have its own password chain based on shuffled order.")
        
    else:
        st.info("ℹ️ Single set mode - only one version will be generated")
        st.session_state.num_combinations = 1
        st.session_state.shuffle_mode = "none"
        st.session_state.keep_first_fixed = False
        st.session_state.keep_last_fixed = False
    
    st.markdown("---")
    
    # ==========================================================================
    # STEP 5: OUTPUT CONFIGURATION
    # ==========================================================================
    
    st.header("Step 5: Output Configuration")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("File Naming")
        
        naming_option = st.selectbox(
            "Naming convention:",
            [
                "Section_1.pdf, Section_2.pdf, ...",
                "Document_1.pdf, Document_2.pdf, ...",
                "Part_1.pdf, Part_2.pdf, ...",
                "Custom prefix"
            ]
        )
        
        if naming_option == "Custom prefix":
            custom_prefix = st.text_input(
                "Enter prefix:",
                placeholder="MyDoc",
                help="e.g., MyDoc_1.pdf"
            )
            file_prefix = custom_prefix if custom_prefix else "Section"
        else:
            file_prefix = naming_option.split("_")[0]
        
        st.info(f"📝 Files: `{file_prefix}_1.pdf`, `{file_prefix}_2.pdf`, ...")
    
    with col2:
        st.subheader("Preview")
        st.markdown("**Sample filenames:**")
        for i in range(min(3, len(st.session_state.sections))):
            lock = "🔒" if passwords.get(i+1) else "🔓"
            st.text(f"{lock} {file_prefix}_{i+1}.pdf")
        if len(st.session_state.sections) > 3:
            st.text(f"... ({len(st.session_state.sections)} total)")
    
    st.markdown("---")
    
    # ==========================================================================
    # STEP 6: GENERATE
    # ==========================================================================
    
    st.header("Step 6: Generate PDFs")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        generate_button = st.button(
            "🚀 Generate PDFs",
            type="primary",
            use_container_width=True
        )
    
    if generate_button:
        # Pre-flight checks
        preflight_errors = []
        
        # Check disk space (at least 100MB)
        try:
            stat = shutil.disk_usage(tempfile.gettempdir())
            free_mb = stat.free / (1024 * 1024)
            if free_mb < 100:
                preflight_errors.append(f"⚠️ Low disk space: {free_mb:.0f}MB available (need at least 100MB)")
        except:
            pass  # Skip if can't check
        
        # Validate all passwords if in sequential mode
        if st.session_state.mode == 'quiz' and st.session_state.encryption_mode == 'sequential':
            if not st.session_state.passwords:
                preflight_errors.append("⚠️ No passwords set for Sequential mode")
        
        # Show preflight errors if any
        if preflight_errors:
            for error in preflight_errors:
                st.warning(error)
            
            if not st.button("⚠️ Continue Anyway"):
                st.stop()
        
        # Start generation with enhanced error handling
        with st.spinner("⏳ Initializing..."):
            try:
                # Create progress indicators
                progress_container = st.container()
                with progress_container:
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    time_estimate = st.empty()
                
                import time
                start_time = time.time()
                
                temp_pdf_dir = tempfile.mkdtemp()
                
                pdf_generator = PDFGenerator()
                pdf_encryptor = PDFEncryptor()
                
                status_text.text("📊 Planning combinations...")
                
                # Generate combinations if enabled
                if st.session_state.combinations_enabled:
                    if st.session_state.mode == 'quiz' and st.session_state.shuffle_mode == "grouped":
                        # Use quiz generator for grouped shuffling
                        all_combinations = QuizGenerator.generate_grouped_combinations(
                            st.session_state.sections,
                            st.session_state.num_combinations,
                            st.session_state.quiz_groups,
                            st.session_state.quiz_shuffle_settings
                        )
                    else:
                        # Use regular combination generator
                        all_combinations = CombinationGenerator.generate_combinations(
                            st.session_state.sections,
                            st.session_state.num_combinations,
                            st.session_state.shuffle_mode,
                            st.session_state.keep_first_fixed,
                            st.session_state.keep_last_fixed
                        )
                else:
                    all_combinations = [st.session_state.sections]
                
                st.session_state.all_combinations = all_combinations
                
                # Store all generated files
                all_generated_files = []
                combination_folders = []
                
                total_pdfs = len(all_combinations) * len(st.session_state.sections)
                current_pdf = 0
                
                status_text.text(f"🎯 Generating {total_pdfs} PDFs across {len(all_combinations)} combination(s)...")
                
                # Track failures for summary
                failed_pdfs = []
                
                # Generate PDFs for each combination
                for combo_idx, combination in enumerate(all_combinations, start=1):
                    # Create folder for this combination
                    if st.session_state.combinations_enabled:
                        combo_name = f"Combo_{combo_idx:02d}"
                        combo_dir = os.path.join(temp_pdf_dir, combo_name)
                        os.makedirs(combo_dir, exist_ok=True)
                        combination_folders.append(combo_dir)
                    else:
                        combo_dir = temp_pdf_dir
                    
                    # Generate PDFs for this combination
                    for idx, section in enumerate(combination):
                        current_pdf += 1
                        progress = current_pdf / total_pdfs
                        progress_bar.progress(progress)
                        
                        # Calculate ETA
                        elapsed = time.time() - start_time
                        if current_pdf > 1:
                            avg_time_per_pdf = elapsed / current_pdf
                            remaining_pdfs = total_pdfs - current_pdf
                            eta_seconds = avg_time_per_pdf * remaining_pdfs
                            time_estimate.text(f"⏱️ Est. remaining: {eta_seconds:.0f}s ({current_pdf}/{total_pdfs} PDFs)")
                        
                        if st.session_state.combinations_enabled:
                            status_text.text(f"📄 Combo {combo_idx}/{len(all_combinations)} - {file_prefix}_{idx+1}.pdf...")
                        else:
                            status_text.text(f"📄 Creating {file_prefix}_{idx+1}.pdf...")
                        
                        pdf_filename = f"{file_prefix}_{idx+1}.pdf"
                        temp_pdf = os.path.join(combo_dir, f"temp_{pdf_filename}")
                        final_pdf = os.path.join(combo_dir, pdf_filename)
                        
                        try:
                            # Generate PDF with error handling
                            pdf_generator.create_pdf(section, temp_pdf)
                            
                            # Handle encryption
                            password = None
                            
                            if st.session_state.encryption_mode == "same":
                                # Same password for all
                                password = passwords.get(1)  # Use the common password
                            
                            elif st.session_state.encryption_mode == "different":
                                # Different password per ORIGINAL section number
                                password = passwords.get(section.original_number)
                            
                            elif st.session_state.encryption_mode == "sequential":
                                # Sequential: first has no password, others based on position
                                if idx == 0:  # First in this combo
                                    password = None
                                else:
                                    # Use password for this POSITION (not original number)
                                    password = passwords.get(idx + 1) if passwords.get(idx + 1) else None
                            
                            elif st.session_state.encryption_mode == "none":
                                password = None
                            
                            # Apply encryption if password exists
                            if password:
                                try:
                                    pdf_encryptor.encrypt_pdf(temp_pdf, final_pdf, password)
                                    os.remove(temp_pdf)
                                except Exception as enc_error:
                                    st.warning(f"⚠️ Encryption failed for {pdf_filename}: {str(enc_error)}")
                                    # Fall back to unencrypted
                                    os.rename(temp_pdf, final_pdf)
                                    failed_pdfs.append((pdf_filename, "encryption failed"))
                            else:
                                os.rename(temp_pdf, final_pdf)
                            
                            all_generated_files.append(final_pdf)
                        
                        except Exception as pdf_error:
                            error_msg = ErrorHandler.handle_pdf_error(pdf_error)
                            st.warning(f"⚠️ Failed to create {pdf_filename}: {error_msg}")
                            failed_pdfs.append((pdf_filename, str(pdf_error)))
                            # Continue with other PDFs
                            continue
                
                # Clear time estimate
                time_estimate.empty()
                
                # Store results
                st.session_state.generated_pdfs = {
                    'files': all_generated_files,
                    'temp_dir': temp_pdf_dir,
                    'prefix': file_prefix,
                    'combinations_enabled': st.session_state.combinations_enabled,
                    'combination_folders': combination_folders if st.session_state.combinations_enabled else None,
                    'num_combinations': len(all_combinations),
                    'generation_time': time.time() - start_time,
                    'failed_pdfs': failed_pdfs
                }
                
                progress_bar.progress(1.0)
                status_text.text("✅ Generation complete!")
                
                # Show summary with timing
                elapsed_time = time.time() - start_time
                
                if st.session_state.combinations_enabled:
                    st.success(f"🎉 Generated {st.session_state.num_combinations} combinations with {len(st.session_state.sections)} PDFs each in {elapsed_time:.1f}s!")
                    st.info(f"📊 Total: {len(all_generated_files)} PDFs created")
                else:
                    st.success(f"🎉 Generated {len(all_generated_files)} PDF(s) in {elapsed_time:.1f}s!")
                
                # Show failures if any
                if failed_pdfs:
                    with st.expander(f"⚠️ {len(failed_pdfs)} PDF(s) had issues"):
                        for filename, error in failed_pdfs:
                            st.text(f"- {filename}: {error}")
                
                # Generate Quiz Mode extras
                if st.session_state.mode == 'quiz':
                    status_text.text("📝 Generating answer keys and checklist...")
                    
                    try:
                        # Generate answer keys
                        st.session_state.answer_keys = QuizGenerator.generate_answer_keys(
                            all_combinations,
                            st.session_state.passwords,
                            "Combo"
                        )
                        
                        # Generate verification checklist
                        st.session_state.verification_checklist = QuizGenerator.generate_verification_checklist(
                            all_combinations,
                            st.session_state.passwords,
                            st.session_state.encryption_mode,
                            "Combo"
                        )
                        
                        st.success("✅ Answer keys and verification checklist generated!")
                    except Exception as quiz_error:
                        st.warning(f"⚠️ Could not generate quiz extras: {str(quiz_error)}")
                
                # Show generation statistics
                with st.expander("📊 Generation Statistics"):
                    st.metric("Total PDFs", len(all_generated_files))
                    st.metric("Total Time", f"{elapsed_time:.1f}s")
                    st.metric("Avg Time/PDF", f"{elapsed_time/len(all_generated_files):.2f}s")
                    if failed_pdfs:
                        st.metric("Failed", len(failed_pdfs))
                
            except Exception as e:
                error_msg = ErrorHandler.handle_file_error(e, "PDF generation")
                st.error(f"❌ {error_msg}")
                
                with st.expander("🔍 Technical Details"):
                    import traceback
                    st.code(traceback.format_exc())
                
                st.info("💡 Try: Close any open PDFs, check disk space, or try with fewer combinations.")
                st.stop()
    
    # ==========================================================================
    # STEP 7: DOWNLOAD
    # ==========================================================================
    
    if st.session_state.generated_pdfs:
        st.markdown("---")
        st.header("Step 7: Download")
        
        # Check if combinations were generated
        if st.session_state.generated_pdfs.get('combinations_enabled'):
            # Multiple combinations - show different download UI
            st.subheader("📦 Download Combinations")
            
            num_combos = st.session_state.generated_pdfs['num_combinations']
            combo_folders = st.session_state.generated_pdfs['combination_folders']
            
            st.info(f"✅ Generated {num_combos} combinations with {len(st.session_state.sections)} PDFs each")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### 📦 Download All Combinations")
                st.markdown(f"One ZIP containing all {num_combos} combination folders")
                
                # Create master ZIP with all combinations
                master_zip = os.path.join(st.session_state.generated_pdfs['temp_dir'], "All_Combinations.zip")
                
                with zipfile.ZipFile(master_zip, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for combo_folder in combo_folders:
                        combo_name = os.path.basename(combo_folder)
                        for file in os.listdir(combo_folder):
                            file_path = os.path.join(combo_folder, file)
                            zip_file.write(file_path, os.path.join(combo_name, file))
                
                with open(master_zip, 'rb') as f:
                    master_zip_bytes = f.read()
                
                st.download_button(
                    label=f"⬇️ Download All ({num_combos} Combos)",
                    data=master_zip_bytes,
                    file_name="All_Combinations.zip",
                    mime="application/zip",
                    use_container_width=True
                )
            
            with col2:
                st.markdown("### 📄 Individual Combinations")
                st.markdown(f"Download specific combinations separately")
                
                # Show scrollable list of individual combination downloads
                for combo_folder in combo_folders:
                    combo_name = os.path.basename(combo_folder)
                    
                    # Create ZIP for this combination
                    combo_zip = os.path.join(combo_folder, f"{combo_name}.zip")
                    
                    with zipfile.ZipFile(combo_zip, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                        for file in os.listdir(combo_folder):
                            if file.endswith('.pdf'):
                                file_path = os.path.join(combo_folder, file)
                                zip_file.write(file_path, file)
                    
                    with open(combo_zip, 'rb') as f:
                        combo_zip_bytes = f.read()
                    
                    st.download_button(
                        label=f"⬇️ {combo_name} ({len(st.session_state.sections)} PDFs)",
                        data=combo_zip_bytes,
                        file_name=f"{combo_name}.zip",
                        mime="application/zip",
                        key=f"combo_dl_{combo_name}"
                    )
            
            # Show combination details
            with st.expander("📋 Combination Details"):
                st.markdown("**How sections are ordered in each combination:**")
                
                for idx, combo in enumerate(st.session_state.all_combinations, start=1):
                    order = [f"Q{s.original_number}" for s in combo]
                    st.text(f"Combo_{idx:02d}: {' → '.join(order)}")
            
            # Quiz Mode Extras
            if st.session_state.mode == 'quiz':
                st.markdown("---")
                st.subheader("🎓 Quiz Mode Extras")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("### 📝 Answer Keys")
                    
                    # Generate answer keys text file
                    answer_keys_text = []
                    answer_keys_text.append("=" * 70)
                    answer_keys_text.append("QUIZ ANSWER KEYS")
                    answer_keys_text.append("=" * 70)
                    answer_keys_text.append("")
                    
                    for combo_name, answers in st.session_state.answer_keys.items():
                        answer_keys_text.append(f"\n{combo_name}")
                        answer_keys_text.append("-" * 70)
                        
                        for ans_data in answers:
                            answer_keys_text.append(
                                f"Position {ans_data['position']}: "
                                f"Q{ans_data['original_question']} → {ans_data['answer']}"
                            )
                    
                    answer_keys_content = "\n".join(answer_keys_text)
                    
                    st.download_button(
                        label="📝 Download Answer Keys",
                        data=answer_keys_content,
                        file_name="Quiz_Answer_Keys.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                    
                    st.caption(f"Contains answers for all {len(st.session_state.answer_keys)} combinations")
                
                with col2:
                    st.markdown("### ✅ Verification Checklist")
                    
                    st.download_button(
                        label="✅ Download Checklist",
                        data=st.session_state.verification_checklist,
                        file_name="Quiz_Verification_Checklist.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                    
                    st.caption("Use this to test all combinations before distribution")
                
                # Preview answer keys
                with st.expander("👁️ Preview Answer Keys"):
                    st.text(answer_keys_content[:1000] + "..." if len(answer_keys_content) > 1000 else answer_keys_content)
                
                # Preview checklist
                with st.expander("👁️ Preview Verification Checklist"):
                    st.text(st.session_state.verification_checklist[:1000] + "..." if len(st.session_state.verification_checklist) > 1000 else st.session_state.verification_checklist)
        
        else:
            # Single set - original download UI
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📦 Download All")
                
                zip_path = create_download_zip(
                    st.session_state.generated_pdfs['files'],
                    f"{file_prefix}_PDFs.zip"
                )
                
                with open(zip_path, 'rb') as f:
                    zip_bytes = f.read()
                
                st.download_button(
                    label="⬇️ Download ZIP",
                    data=zip_bytes,
                    file_name=f"{file_prefix}_PDFs.zip",
                    mime="application/zip",
                    use_container_width=True
                )
                
                if st.session_state.encryption_mode != "none":
                    st.info("🔐 PDFs are password-protected")
            
            with col2:
                st.subheader("📄 Individual Files")
                
                st.markdown(f"**{len(st.session_state.generated_pdfs['files'])} files:**")
                
                for pdf_path in st.session_state.generated_pdfs['files']:
                    filename = os.path.basename(pdf_path)
                    
                    with open(pdf_path, 'rb') as f:
                        pdf_bytes = f.read()
                    
                    # Determine lock icon
                    if st.session_state.encryption_mode == "none":
                        lock_icon = "🔓"
                    elif st.session_state.encryption_mode == "sequential":
                        # First file is unlocked
                        section_num = int(filename.split('_')[-1].replace('.pdf', ''))
                        lock_icon = "🔓" if section_num == 1 else "🔒"
                    else:
                        lock_icon = "🔒"
                    
                    st.download_button(
                        label=f"{lock_icon} {filename}",
                        data=pdf_bytes,
                        file_name=filename,
                        mime="application/pdf",
                        key=f"dl_{filename}"
                    )
        
        st.markdown("---")
        
        # Password reference
        if st.session_state.encryption_mode in ["different", "sequential"]:
            with st.expander("🔑 Password Reference (for your records)"):
                st.warning("⚠️ Keep this information secure!")
                
                if st.session_state.combinations_enabled and st.session_state.encryption_mode == "sequential":
                    st.info("ℹ️ Sequential mode with combinations: Password chain differs per combination based on shuffled order")
                
                for section in st.session_state.sections:
                    pwd = passwords.get(section.number)
                    if pwd:
                        st.text(f"Section {section.number}: {pwd}")
                    else:
                        st.text(f"Section {section.number}: No password")
        
        # Reset
        if st.button("🔄 Start Over", type="secondary"):
            cleanup_temp_files(st.session_state.generated_pdfs['temp_dir'])
            
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            
            st.rerun()


# =============================================================================
# ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    main()